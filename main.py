from tkinter import *
from tkinter import ttk
from tkinter.filedialog import askopenfilename, asksaveasfilename
import os, sv_ttk
from docx import Document

# Create a window
win = Tk()
win.title("IMI-StoryWriter")
win.geometry("419x595")  # Set the geometry to A5 size in pixels
document = Document()

# Colorscbhemes
def get_system_preference():
    if os.name == 'posix':  # macOS or Linux
        # Check if Apple's NSUserDefaults exist
        if os.system('defaults read -g AppleInterfaceStyle > /dev/null 2>&1') == 0:
            return 'dark'  # Dark mode is enabled
        else:
            return 'light'  # Light mode is enabled
    elif os.name == 'nt':  # Windows
        # Check if the Registry key exists
        try:
            import winreg
            reg_path = r'SOFTWARE\Microsoft\Windows\CurrentVersion\Themes\Personalize'
            reg_key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path)
            value, _ = winreg.QueryValueEx(reg_key, 'AppsUseLightTheme')
            if value == 0:
                return 'dark'  # Dark mode is enabled
            else:
                return 'light'  # Light mode is enabled
        except Exception:
            pass
    return 'unknown'  # Unable to determine the system preference

# Usage example
preference = get_system_preference()
print(f'System preference: {preference} mode')

if preference == "dark":
    sv_ttk.set_theme("dark")

    bg = "#1C1C1C"
    accent = "#57C8FF"
    surface = "#2A2A2A"
    white = "white"
    subcolor = '#9FB6C6'

else:
    sv_ttk.set_theme('light')

    bg = "#FAFAFA"
    accent = "#0560B6"
    surface = '#cfdbe3'
    white = '#1C1C1C'
    subcolor = '#4d5961'


SIZE = 14

# Create main frame
root = ttk.Frame()
root.pack(fill='both', expand=1)

# Create a topbar
tabbar = Frame(root, bg=surface)
tabbar.pack(side='top', fill='x')

# Bottom Bar
bar = Frame(root, bg=surface)
bar.pack(side='bottom', fill='x')

home_frame = Frame(root, bg=bg)

# Home
def fun_home():
    row_export.config(bg=surface)
    row_open.config(bg=surface)
    row_home.config(bg=accent)

# Export
def fun_export():
    global document

    row_home.config(bg=surface)
    row_open.config(bg=surface)
    row_export.config(bg=accent)

    # Clear the document
    document = Document()

    first_line = text.get("1.0", "1.end")
    print("First line:", first_line)
    document.add_heading(first_line, 0)

    # Get the second line
    second_line = text.get("2.0", "2.end")
    document.add_heading(f"{second_line}\n", level=4)

    # Get the remaining text
    remaining_text = text.get("3.0", "end")
    paragraphs = remaining_text.split("\n")
    for paragraph in paragraphs:
        if paragraph.strip():  # Skip empty lines
            print(paragraph)

            if paragraph.startswith("#####"):
                document.add_heading(paragraph.strip(), level=4)
            elif paragraph.startswith("####"):
                document.add_heading(paragraph.strip(), level=3)
            elif paragraph.startswith("###"):
                document.add_heading(paragraph.strip(), level=2)
            elif paragraph.startswith("##"):
                document.add_heading(paragraph.strip(), level=1)
            elif paragraph.startswith("#"):
                document.add_heading(paragraph.strip(), level=0)
            else:
                if paragraph != "":
                    document.add_paragraph(paragraph)

    file_path = asksaveasfilename(defaultextension=".docx")
    document.save(file_path)
    os.startfile(file_path)

    # Clear the document
    document = Document()

# Open
def fun_open():
    global file_path

    row_home.config(bg=surface)
    row_open.config(bg=accent)
    row_export.config(bg=surface)

    file_path = askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        # Clear the textbox
        text.delete("1.0", "end")

        # Read the contents of the selected Word document
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            text.insert("end", paragraph.text + "\n\n")

        # Filepath
        root.config(title=f'IMI-StoryWriter | {file_path}')

# Home
Label(tabbar, text=" ", bg=surface).grid(row=0, column=0)

# Tab home
tab_home = Label(tabbar, text='Home', bg=surface, fg=white, font=("Product Sans", 12))
tab_home.grid(row=0, column=1)

row_home = Frame(tabbar, bg=surface, width=50, height=2)
row_home.grid(row=1, column=1, columnspan=1)

tab_home.bind("<Button-1>", lambda e: fun_home())

# tab Export
tab_export = Label(tabbar, text='Export', bg=surface, fg=white, font=("Product Sans", 12))
tab_export.grid(row=0, column=2)

row_export = Frame(tabbar, bg=surface, width=50, height=2)
row_export.grid(row=1, column=2, columnspan=1)

tab_export.bind("<Button-1>", lambda e: fun_export())

# Tab Open
tab_open = Label(tabbar, text='Open', bg=surface, fg=white, font=("Product Sans", 12))
tab_open.grid(row=0, column=4)

row_open = Frame(tabbar, bg=surface, width=50, height=2)
row_open.grid(row=1, column=4, columnspan=1)

tab_open.bind("<Button-1>", lambda e: fun_open())

# Textbox
text = Text(root, bg=bg, fg=white, selectbackground=accent, selectforeground=bg, font=("Roboto Regular", SIZE), bd=0,
            wrap=WORD, undo=True)
text.pack(fill='both', expand=1, pady=6, padx=6)

# Word Count
def update_word_count(event=None):
    text_content = text.get("1.0", "end-1c")
    word_count = len(text_content.split())
    letter_count = len(text_content.replace(" ", ""))
    status_label.config(text=f"Word Count: {word_count}")
    letter_count_label.config(text=f"Letter Count: {letter_count}")

    # Add tag to first line
    text.tag_add("tag_first_line", "1.0", "1.end")
    text.tag_configure("tag_first_line", font=("Product Sans", 26), foreground=accent, justify='center')

    text.tag_add("tag_second_line", "2.0", "2.end")
    text.tag_configure("tag_second_line", font=("Product Sans", 18), foreground=subcolor, justify='center')

    # Iterate over each line
    lines = text_content.split("\n")
    for line_num, line in enumerate(lines, start=1):
        if line.startswith("#"):
            tag_name = f"tag_heading{line.count('#')}"
            font_size = 24 - line.count('#') * 2  # Adjust the font size based on heading level
            text.tag_configure(tag_name, font=("Product Sans", font_size), foreground=accent)
            start_index = f"{line_num}.0"
            end_index = f"{line_num}.end"
            text.tag_add(tag_name, start_index, end_index)

            # Add the corresponding heading to the document
            heading_level = line.count('#')
            heading_text = line.strip('#').strip()
            document.add_heading(heading_text, level=heading_level)

# Bind the update_word_count function to appropriate events
text.bind("<<Modified>>", update_word_count)
text.bind("<KeyRelease>", update_word_count)

# Status Bar
status_label = Label(bar, text="Word Count: 0", bg=surface, fg=white, anchor="e", padx=10)
status_label.pack(side="right")

# Letter Count
letter_count_label = Label(bar, text="Letter Count: 0", bg=surface, fg=white, anchor="e", padx=5)
letter_count_label.pack(side="right")

# Execute
if __name__ == "__main__":
    fun_home()
    win.mainloop()
